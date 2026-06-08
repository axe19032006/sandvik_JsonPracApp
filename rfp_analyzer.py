import streamlit as st          # the web app framework
from docx import Document       # reads .docx Word files
import io                       # lets us treat raw bytes as a file

# ============================================================
# PAGE SETUP
# ============================================================
st.set_page_config(page_title="Sandvik RFP Analyser", layout="wide")
st.title("Sandvik RFP Analyser")

# ============================================================
# INITIALIZE SESSION STATE
# ============================================================
# These are all the things our app will need to remember.
# We set them to None now; they get filled in as the user works.
for key in ["rfp_text", "parameters", "proposals", "results"]:
    if key not in st.session_state:
        st.session_state[key] = None

# ============================================================
# HELPER FUNCTION: read a .docx file and return its text
# ============================================================
def read_docx_bytes(file_bytes: bytes) -> str:
    # file_bytes is the raw binary data of the Word file
    # io.BytesIO() wraps those bytes so Document() can open them
    # (Document() normally opens a file path, but BytesIO lets
    #  us give it bytes directly without saving to disk)
    doc = Document(io.BytesIO(file_bytes))

    # doc.paragraphs is a list of every paragraph in the document
    # Each paragraph has a .text property with the actual words
    # We skip empty paragraphs with "if p.text.strip()"
    # We join everything with newlines to get one big text string
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

# ============================================================
# TWO TABS
# ============================================================
tab1, tab2 = st.tabs(["1. Upload RFP", "2. Analyse Proposals"])

# ============================================================
# TAB 1: UPLOAD AND READ THE RFP
# ============================================================
with tab1:
    st.header("Upload the RFP")
    st.write("Upload the Sandvik RFP document. The app will read its text.")

    # File uploader — only accepts .docx files
    rfp_file = st.file_uploader(
        "Upload RFP (.docx)",
        type=["docx"],          # only Word files allowed
        accept_multiple_files=False,  # only one RFP at a time
    )

    # When a file is uploaded, read it immediately
    if rfp_file:
        # rfp_file.read() gives us the raw bytes of the uploaded file
        raw_bytes = rfp_file.read()

        # Pass those bytes to our helper function
        rfp_text = read_docx_bytes(raw_bytes)

        # Save the text to session state so Tab 2 can use it later
        st.session_state.rfp_text = rfp_text

        st.success(f"RFP uploaded: {rfp_file.name}")

    # If we have RFP text stored, show it
    if st.session_state.rfp_text:
        st.subheader("Extracted text (preview)")
        # st.text_area shows the text in a scrollable box
        # value= is the content; height= controls how tall the box is
        st.text_area(
            label="RFP text",
            value=st.session_state.rfp_text,
            height=400,
        )
        # Show how many characters were extracted
        char_count = len(st.session_state.rfp_text)
        st.caption(f"Total characters extracted: {char_count:,}")

# ============================================================
# TAB 2: PLACEHOLDER FOR NOW
# ============================================================
with tab2:
    st.header("Analyse Proposals")